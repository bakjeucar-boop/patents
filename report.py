from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.models import PatentAnalysis, PatentDocument


def save_analysis(output_dir: Path, document: PatentDocument, analysis: PatentAnalysis) -> Path:
    stem = document.source.label.replace("/", "_").replace("\\", "_").replace(":", "_")[:80] or "patent"
    path = output_dir / f"{stem}.analysis.json"
    payload = {
        "document": document.model_dump(),
        "analysis": analysis.model_dump(),
    }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def _title(document: PatentDocument, analysis: PatentAnalysis) -> str:
    if document.publication_number and document.publication_number in analysis.title:
        return analysis.title.replace(document.publication_number, "").strip(" -:|") or analysis.title
    return analysis.title


def build_markdown_report(documents: list[PatentDocument], analyses: list[PatentAnalysis]) -> str:
    lines = ["# 특허 검토 보고서", ""]
    for idx, (document, analysis) in enumerate(zip(documents, analyses), start=1):
        lines.extend(
            [
                f"## P{idx}. {_title(document, analysis)}",
                "",
                f"- 특허번호: {document.publication_number or '-'}",
                f"- 출처: {document.source.label}",
                f"- 신뢰도: {analysis.confidence}",
                "",
                "### 요약",
                analysis.one_line_summary,
                "",
                "### 핵심 요약",
                analysis.simple_explanation,
                "",
                "### 왜 특허가 되었는가",
                analysis.why_patentable,
                "",
                "### 해결하려는 문제",
                analysis.problem_to_solve,
                "",
                "### 기술 차별점",
            ]
        )
        lines.extend(f"- {item}" for item in analysis.differentiators)
        lines.extend(["", "### 청구항 구성요소"])
        for element in analysis.key_claim_elements:
            lines.append(f"- {element.label}. {element.plain_explanation} ({element.text})")
        lines.extend(["", "### 회피설계 검토 포인트"])
        lines.extend(f"- {item}" for item in analysis.design_around_points)
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def build_docx_report(documents: list[PatentDocument], analyses: list[PatentAnalysis]) -> bytes:
    report = Document()
    section = report.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = report.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color, before, after in [
        ("Title", 24, "0B2545", 0, 12),
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = report.add_paragraph("특허 검토 보고서", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle = report.add_paragraph(f"분석 대상 {len(analyses)}건 | 엔지니어링 기술 검토용")
    subtitle.runs[0].font.color.rgb = RGBColor(90, 100, 115)

    report.add_heading("분석 대상 요약", level=1)
    summary_table = report.add_table(rows=1, cols=4)
    summary_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    summary_table.autofit = False
    widths = [Inches(0.55), Inches(1.25), Inches(2.35), Inches(2.35)]
    headers = ["ID", "특허번호", "제목", "요약"]
    for idx, (cell, header, width) in enumerate(zip(summary_table.rows[0].cells, headers, widths)):
        cell.width = width
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade_cell(cell, "F2F4F7")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for idx, (document, analysis) in enumerate(zip(documents, analyses), start=1):
        cells = summary_table.add_row().cells
        values = [f"P{idx}", document.publication_number or "-", _title(document, analysis), analysis.one_line_summary]
        for cell, value, width in zip(cells, values, widths):
            cell.width = width
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    for idx, (document, analysis) in enumerate(zip(documents, analyses), start=1):
        report.add_heading(f"P{idx}. {_title(document, analysis)}", level=1)
        meta = report.add_paragraph()
        meta.add_run("특허번호  ").bold = True
        meta.add_run(document.publication_number or "-")
        meta.add_run("    신뢰도  ").bold = True
        meta.add_run(analysis.confidence)

        _add_section(report, "요약", analysis.one_line_summary)
        _add_section(report, "핵심 요약", analysis.simple_explanation)
        _add_section(report, "왜 특허가 되었는가", analysis.why_patentable)
        _add_section(report, "해결하려는 문제", analysis.problem_to_solve)
        _add_bullets(report, "기술 차별점", analysis.differentiators)
        claim_items = [f"{element.label}. {element.plain_explanation} ({element.text})" for element in analysis.key_claim_elements]
        _add_bullets(report, "청구항 구성요소", claim_items)
        _add_bullets(report, "회피설계 검토 포인트", analysis.design_around_points)

    footer = section.footer.paragraphs[0]
    footer.text = "특허 검토 Agent | 기술 검토 보조 자료"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(110, 120, 130)

    buffer = BytesIO()
    report.save(buffer)
    return buffer.getvalue()


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _add_section(report: Document, title: str, body: str) -> None:
    report.add_heading(title, level=2)
    report.add_paragraph(body or "-")


def _add_bullets(report: Document, title: str, items: list[str]) -> None:
    report.add_heading(title, level=2)
    if not items:
        report.add_paragraph("-")
        return
    for item in items:
        paragraph = report.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)

